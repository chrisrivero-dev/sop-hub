print("🔥 REAL APP.PY LOADED FROM:", __file__)
print("🔥 RUNNING APP.PY FROM:", __file__)

# ======================================================
# IMPORTS
# ======================================================
from flask import Flask, render_template, request, jsonify, redirect, url_for
from flask_sqlalchemy import SQLAlchemy
from flask_migrate import Migrate
from datetime import datetime
from services.workspace_service import copy_reference_to_workspace
from services.excel_preview_service import build_excel_preview
from services.word_preview_service import build_word_preview
from routes.file_summary_routes import file_summary_bp
from routes.topic_summary_routes import topic_summary_bp
import docx
import os
import re
from collections import defaultdict
from sqlalchemy.exc import IntegrityError, SQLAlchemyError
from sqlalchemy import or_, text
import threading
import sys
import webbrowser
import time
import pandas as pd
from flask import send_file, abort, request
import subprocess
import tempfile


BASE_DIR = os.path.abspath(os.path.dirname(__file__))

EXCEL_FILENAME = "ParcelingLogCRTEMP.xlsx"

# In-memory cache: avoid re-reading Excel on every search request.
_excel_cache = {"mtime": None, "rows": []}


def _load_excel_rows():
    """Parse ParcelingLogCRTEMP.xlsx into a list of row dicts. Cached by mtime."""
    excel_path = os.path.join(BASE_DIR, EXCEL_FILENAME)

    if not os.path.exists(excel_path):
        return []

    try:
        mtime = os.path.getmtime(excel_path)
    except OSError:
        return []

    if _excel_cache["mtime"] == mtime and _excel_cache["rows"] is not None:
        return _excel_cache["rows"]

    rows = []
    try:
        xl = pd.ExcelFile(excel_path)
        target_sheets = ["PARCELING", "DOC SCREENING"]

        for sheet in target_sheets:
            if sheet not in xl.sheet_names:
                continue

            df = xl.parse(sheet)
            df = df.fillna("")

            for _, row in df.iterrows():
                drn = str(row.get("DRN", "")).strip()
                action = str(row.get("Action", "")).strip()

                if not drn.isdigit() or not action:
                    continue

                remarks = str(row.get("Remarks", "")).strip()
                apn = str(row.get("APN", "")).strip()
                date = str(row.get("Date", "")).strip()
                notes = str(row.get("Notes", "")).strip()
                example_file_raw = str(row.get("Example_File", "")).strip()

                example_file_path = ""
                if example_file_raw:
                    example_file_path = os.path.join(
                        r"Q:\Mapping Reference Workspace",
                        example_file_raw
                    )

                rows.append({
                    "drn": drn,
                    "action": action,
                    "remarks": remarks,
                    "apn": apn,
                    "date": date,
                    "notes": notes,
                    "example_file": example_file_path,
                    "sheet": sheet,
                    "_combined": " ".join([drn, action, remarks, apn, date, notes]).lower(),
                })

        _excel_cache["mtime"] = mtime
        _excel_cache["rows"] = rows

    except Exception as e:
        print("Excel read error:", e)

    return rows


def load_excel_examples(query):
    rows = _load_excel_rows()

    if not query:
        return []

    q = query.lower()
    matched = [r for r in rows if q in r["_combined"]]

    return [{k: v for k, v in r.items() if k != "_combined"} for r in matched[:10]]

template_dir = os.path.join(BASE_DIR, "templates")
static_dir   = os.path.join(BASE_DIR, "static")

app = Flask(
    __name__,
    static_folder=static_dir,
    template_folder=template_dir
)
app.register_blueprint(file_summary_bp)
app.register_blueprint(topic_summary_bp)

# ======================================================
# PREVIEW SYSTEM (THUMBNAIL-BASED)
# ======================================================

def get_preview_path(original_path):
    base, _ = os.path.splitext(original_path)

    for ext in [".png", ".jpg", ".jpeg"]:
        candidate = base + ext
        if os.path.exists(candidate):
            return candidate

    return None


@app.route("/preview")
def preview_file():
    path = request.args.get("path")

    if not path:
        return "Missing path", 400

    path = os.path.normpath(path)

    if not os.path.exists(path):
        return "File not found", 404

    ext = os.path.splitext(path)[1].lower()

    # Direct image preview
    if ext in [".png", ".jpg", ".jpeg", ".gif", ".webp"]:
        return send_file(path)

    # Direct PDF preview
    if ext == ".pdf":
        return send_file(path)

    # Office / other files: look for generated companion preview
    base, _ = os.path.splitext(path)

    for preview_ext in [".png", ".jpg", ".jpeg"]:
        preview_path = base + preview_ext
        if os.path.exists(preview_path):
            return send_file(preview_path)

    return "No preview available", 404

# ======================================================
# CONSTANTS
# ======================================================

Q_DRIVE_ROOTS = [
    r"Q:\WORD",
    r"Q:\Excel Files",
    r"Q:\Mapping Reference Workspace"
]

# ======================================================
# FLASK SETUP — SUPPORTS BOTH DEV AND ONE-FILE EXE
# ======================================================
if hasattr(sys, "_MEIPASS"):
    # EXE MODE — resources extracted into temp folder
    BASE_DIR = sys._MEIPASS
else:
    # DEV MODE — normal folder with app.py
    BASE_DIR = os.path.abspath(os.path.dirname(__file__))

template_dir = os.path.join(BASE_DIR, "templates")
static_dir   = os.path.join(BASE_DIR, "static")



# ======================================================
# DATABASE LOCATION — LOCALAPPDATA IN EXE MODE
# ======================================================
if getattr(sys, "frozen", False):
    # When running from EXE, store DB safely here:
    root = os.path.join(os.environ["LOCALAPPDATA"], "OC_SOP_Hub")
else:
    # Dev mode stores DB next to app.py
    root = os.path.dirname(__file__)

INSTANCE_DIR = os.path.join(root, "instance")
os.makedirs(INSTANCE_DIR, exist_ok=True)

DB_PATH = os.path.join(INSTANCE_DIR, "sop.db")
app.config["SQLALCHEMY_DATABASE_URI"] = f"sqlite:///{DB_PATH}"

db = SQLAlchemy(app)
migrate = Migrate(app, db)



# ======================================================
# DATABASE MODELS
# ======================================================

class Reference(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    file_name = db.Column(db.String(255), nullable=False)
    file_path = db.Column(db.String(600), nullable=False, unique=True)
    folder = db.Column(db.String(255))

    display_name = db.Column(db.String(255))
    group = db.Column(db.String(100))

    notes = db.Column(db.Text, default="")
    pinned = db.Column(db.Boolean, default=False)
    added_at = db.Column(db.DateTime, default=datetime.utcnow)

    file_type = db.Column(db.String(20))


class Recent(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    ref_id = db.Column(db.Integer)
    file_name = db.Column(db.String(255), nullable=False)
    file_path = db.Column(db.String(600), nullable=False)
    folder = db.Column(db.String(255))
    opened_at = db.Column(db.DateTime, default=datetime.utcnow)
    notes = db.Column(db.Text, default="")
    file_type = db.Column(db.String(20))


class CustomGroup(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    name = db.Column(db.String(255), unique=True, nullable=False)
    created_at = db.Column(db.DateTime, default=datetime.utcnow)


class FileIndex(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    file_path = db.Column(db.String(600), unique=True, nullable=False)
    file_name = db.Column(db.String(255), nullable=False)
    folder = db.Column(db.String(255))
    ext = db.Column(db.String(20))
    mtime = db.Column(db.DateTime, nullable=False)
    content = db.Column(db.Text, default="")
    snippet = db.Column(db.Text, default="")
    indexed_at = db.Column(db.DateTime, default=datetime.utcnow)


# Create DB tables
with app.app_context():
    db.create_all()
    # ======================================================
# CREATE BUILT-IN DEFAULT GROUPS (ON FIRST RUN)
# ======================================================
BUILT_IN_GROUPS = [
    "Ungrouped",
    "Procedures",
    "Cut Sheets",
    "Tracts",
    "Parcel Maps",
    "General"
]

with app.app_context():
    for g in BUILT_IN_GROUPS:
        existing = CustomGroup.query.filter_by(name=g).first()
        if not existing:
            db.session.add(CustomGroup(name=g))
    db.session.commit()



# ======================================================
# DESKTOP SHORTCUT CREATION
# ======================================================

def create_shortcut_once():
    """Create a desktop shortcut on first EXE run."""
    if not getattr(sys, "frozen", False):
        return  # Only for EXE mode

    desktop = os.path.join(os.path.expanduser("~"), "Desktop")
    shortcut = os.path.join(desktop, "OC SOP Hub.lnk")

    if os.path.exists(shortcut):
        return  # Already installed

    try:
        import winshell  # type: ignore
        from win32com.client import Dispatch  # type: ignore

        shell = Dispatch("WScript.Shell")
        exe_path = sys.executable

        sc = shell.CreateShortCut(shortcut)
        sc.Targetpath = exe_path
        sc.WorkingDirectory = os.path.dirname(exe_path)
        sc.IconLocation = exe_path
        sc.save()

        print("📌 Desktop shortcut created.")
    except Exception as e:
        print("Shortcut creation error:", e)


# ======================================================
# FILE TEXT EXTRACTORS
# ======================================================

def extract_docx(filepath):
    text = []
    try:
        doc = docx.Document(filepath)

        for p in doc.paragraphs:
            t = p.text.strip()
            if t:
                text.append(t)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    t = cell.text.strip()
                    if t:
                        text.append(t)

        return "\n".join(text)

    except Exception as e:
        print("DOCX Extract Error:", filepath, e)
        return ""


def extract_pdf(filepath):
    try:
        from pypdf import PdfReader  # type: ignore
        reader = PdfReader(filepath)
        text = ""
        for page in reader.pages:
            text += page.extract_text() or ""
        return text
    except Exception as e:
        print("PDF Extract Error:", filepath, e)
        return ""


def extract_text(filepath):
    ext = filepath.lower().split(".")[-1]

    if ext == "txt":
        try:
            with open(filepath, "r", errors="ignore") as f:
                return f.read()
        except Exception:
            return ""

    # ✔ allow .doc files (filename only; no text reading)
    if ext == "doc":
        return ""

    if ext == "docx":
        return extract_docx(filepath)

    return ""  # default



# ======================================================
# RANKING ENGINE + INDEXED SEARCH
# ======================================================

def score_result(item, query_lower):
    name = item["name"].lower()
    folder = (item.get("folder") or "").lower()

    score = 0

    if name == query_lower or name.replace(" ", "") == query_lower.replace(" ", ""):
        score += 1000

    if name.startswith(query_lower):
        score += 600

    if query_lower in name:
        score += 300

    if query_lower in folder:
        score += 50

    if item.get("snippet"):
        score += 100

    return score


def excel_year_from_path(path):
    if not path:
        return 0

    match = re.search(r"(19|20)\d{2}", path)
    if match:
        try:
            return int(match.group(0))
        except Exception:
            return 0
    return 0


def sort_results(ranked, query_lower):
    xls_ext = {"xls", "xlsx", ".xls", ".xlsx"}

    def is_excel(rec):
        ft = (rec.get("file_type") or "").lower()
        return ft in xls_ext

    xls_files = [r for r in ranked if is_excel(r)]
    other_files = [r for r in ranked if not is_excel(r)]

    try:
        xls_files.sort(
            key=lambda r: excel_year_from_path(r.get("full_path", "")),
            reverse=True,
        )
    except Exception as e:
        print("Excel sort error:", e)

    other_files.sort(
        key=lambda r: score_result(r, query_lower),
        reverse=True,
    )

    return xls_files + other_files

def run_search(query, mode="filename"):
    query_lower = query.lower()
    ranked_results = []
    seen = set()

    # Build allowed-root filters (Q:\WORD, Q:\Excel Files, Q:\Mapping Reference Workspace)
    path_filters = [
        FileIndex.file_path.like(base + "%")
        for base in Q_DRIVE_ROOTS
    ]

    # Query only what we need (avoid FileIndex.query.all())
    q = FileIndex.query.filter(or_(*path_filters))

    if mode == "content":
        q = q.filter(FileIndex.content.ilike(f"%{query}%"))
    else:
        q = q.filter(FileIndex.file_name.ilike(f"%{query}%"))

    indexed = q.limit(500).all()

    # 1) INDEXED SEARCH
    if indexed and len(indexed) > 0:
        # Load all Reference rows once — avoids N+1 per-item queries
        ref_map = {
            r.file_path: r.id
            for r in Reference.query.with_entities(
                Reference.file_path,
                Reference.id,
            ).all()
        }

        for item in indexed:
            snippet = ""

            if mode == "content" and item.content:
                idx = item.content.lower().find(query_lower)
                if idx != -1:
                    snippet = item.content[max(0, idx - 60):idx + 60]

            key = (item.file_path or "").lower()
            if not key or key in seen:
                continue
            seen.add(key)

            ref_id = ref_map.get(item.file_path, 0)

            entry = {
                "name": item.file_name,
                "folder": item.folder,
                "full_path": item.file_path,
                "snippet": snippet,
                "is_pinned": ref_id > 0,
                "ref_id": ref_id,
                "file_type": item.ext,
            }

            ranked_results.append(entry)

        ranked_results = sort_results(ranked_results, query_lower)
        return ranked_results

    # 2) FALLBACK SLOW SEARCH
    print("⚠️ No indexed records found — using slow filesystem search")
    return run_slow_search(query, mode)




# ======================================================
# FALLBACK FILE SYSTEM SEARCH
# ======================================================

def run_slow_search(query, mode="filename"):
    query_lower = query.lower()

    ranked_results = []
    seen = set()

    for base_path in Q_DRIVE_ROOTS:
        for root, dirs, files in os.walk(base_path):
            for file in files:
                if file.startswith("~$") or file.startswith("."):
                    continue
                if file.lower() in ("thumbs.db", "desktop.ini"):
                    continue

                full_path = os.path.join(root, file)
                ext = file.lower().split(".")[-1]
                file_lower = file.lower()

                snippet = ""

                if mode == "content" and ext in ["txt", "docx", "pdf"]:
                    try:
                        text = extract_text(full_path).lower()
                        idx = text.find(query_lower)
                        if idx != -1:
                            snippet = text[max(0, idx - 60): idx + 60]
                        else:
                            if query_lower not in file_lower:
                                continue
                    except Exception:
                        continue
                else:
                    if query_lower not in file_lower:
                        continue

                parent_folder = os.path.basename(root)

                entry = {
                    "name": file,
                    "folder": parent_folder,
                    "full_path": full_path,
                    "snippet": snippet,
                    "ref_id": 0,
                    "is_pinned": False,
                    "file_type": ext,
                }

                key = full_path.lower()
                if key in seen:
                    continue
                seen.add(key)

                ranked_results.append(entry)

    ranked_results = sort_results(ranked_results, query_lower)
    return ranked_results

@app.route("/__crash_test")
def __crash_test():
    raise RuntimeError("CRASH TEST — IF YOU SEE THIS, DEBUGGER WORKS")

# ======================================================
# TOGGLE PIN (FINAL, SAFE, IDEMPOTENT)
# ======================================================

@app.route("/preview")
def preview():
    path = request.args.get("path")

    if not path:
        return abort(400)

    if not os.path.exists(path):
        print("[PREVIEW ERROR] File not found:", path)
        return abort(404)

    try:
        return send_file(path)
    except Exception as e:
        print("[PREVIEW ERROR]", e)
        return abort(500)

def is_generated_preview_asset(file_path):
    ext = os.path.splitext(file_path)[1].lower()

    if ext not in {".png", ".jpg", ".jpeg"}:
        return False

    base, _ = os.path.splitext(file_path)

    source_exts = [
        ".doc",
        ".docx",
        ".xls",
        ".xlsx",
        ".pdf",
        ".ppt",
        ".pptx",
    ]

    return any(os.path.exists(base + source_ext) for source_ext in source_exts)

@app.route("/excel-preview")
def excel_preview():
    file_path = request.args.get("path", "").strip()

    result = build_excel_preview(file_path)

    status_code = 200 if result.get("success") else 400

    return jsonify(result), status_code

@app.route("/word-preview")
def word_preview():
    file_path = request.args.get("path", "").strip()

    result = build_word_preview(file_path)

    status_code = 200 if result.get("success") else 400

    return jsonify(result), status_code

@app.route("/toggle-pin", methods=["POST"])
def toggle_pin():
    data = request.get_json() or {}

    file_path = data.get("file_path", "").strip()
    file_name = data.get("file_name", "").strip()
    folder = data.get("folder", "").strip()

    if not file_name or file_name.startswith("~$") or not file_path:
        return jsonify({"success": False, "ignored": True}), 200

    file_path = os.path.normpath(file_path)

    try:
        ref = Reference.query.filter_by(file_path=file_path).first()

        # Existing record: always allow toggle, including old pinned preview assets.
        if ref:
            ref.pinned = not ref.pinned
            db.session.commit()

            return jsonify({
                "success": True,
                "pinned": ref.pinned,
                "ref_id": ref.id
            }), 200

        # New record: block generated companion preview images only.
        # Standalone image references are allowed.
        if is_generated_preview_asset(file_path):
            return jsonify({
                "success": False,
                "ignored": True,
                "reason": "generated_preview_asset"
            }), 200

        ref = Reference(
            file_name=file_name,
            file_path=file_path,
            folder=folder or "Ungrouped",
            pinned=True,
            group="Ungrouped",
        )

        db.session.add(ref)
        db.session.commit()

        return jsonify({
            "success": True,
            "pinned": True,
            "ref_id": ref.id
        }), 200

    except SQLAlchemyError as e:
        db.session.rollback()
        print("💥 toggle-pin ERROR:", e)
        return jsonify({"success": False, "error": str(e)}), 500

# ======================================================
# INDEXER
# ======================================================

def extract_text_for_indexer(path):
    try:
        _, ext = os.path.splitext(path.lower())
        if ext == ".txt":
            with open(path, "r", encoding="utf-8", errors="ignore") as f:
                return f.read()
    except Exception:
        pass
    return ""


def index_qdrive_files():
    print("⚡ REINDEX STARTED")

    for root in Q_DRIVE_ROOTS:
        print("📁 Scanning:", root)

        for dirpath, dirnames, filenames in os.walk(root):
            for fname in filenames:
                if fname.startswith("~$") or fname.startswith("."):
                    continue
                if fname.lower() in ("thumbs.db", "desktop.ini"):
                    continue

                full_path = os.path.join(dirpath, fname)
                stat = os.stat(full_path)

                mtime = datetime.fromtimestamp(stat.st_mtime)
                _, ext = os.path.splitext(fname)
                ext = ext.lower()

                              # Only index selected file types
                # Only index selected file types
                if ext not in [
                    ".txt",
                    ".doc", ".docx",
                    ".pdf",
                    ".xlsx", ".xls",
                    ".pptx",
                    ".png", ".jpg", ".jpeg",
                ]:
                    continue

                # Skip companion preview images (same basename as a source file)
                if is_generated_preview_asset(full_path):
                    continue




                text = ""
                try:
                    text = extract_text_for_indexer(full_path)
                except Exception:
                    pass

                existing = FileIndex.query.filter_by(file_path=full_path).first()

                if existing:
                    existing.file_name = fname
                    existing.folder = os.path.basename(dirpath)
                    existing.ext = ext
                    existing.mtime = mtime
                    existing.content = text
                    existing.snippet = text[:400]
                    existing.indexed_at = datetime.utcnow()
                else:
                    rec = FileIndex(
                        file_path=full_path,
                        file_name=fname,
                        folder=os.path.basename(dirpath),
                        ext=ext,
                        mtime=mtime,
                        content=text,
                        snippet=text[:400],
                    )
                    db.session.add(rec)

    db.session.commit()
    print("✅ REINDEX COMPLETE")


@app.route("/admin/reindex-qdrive")
def admin_reindex_qdrive():
    index_qdrive_files()
    return "Q:\\ index refreshed."


# ======================================================
# AUTO-INDEX ON FIRST RUN (EXE ONLY)
# ======================================================

def should_index():
    """Return True if FileIndex table is empty (first run)."""
    try:
        return FileIndex.query.count() == 0
    except Exception:
        return True


def auto_index_if_needed():
    """
    DISABLED: Indexing is now manual.

    Techs must click "Rebuild File Index" in the Search Hub.
    This keeps first-time behavior consistent (yellow banner).
    """
    print("🔕 auto_index_if_needed(): skipped (manual indexing only).")
    return



# ======================================================
# BROWSER LAUNCH — ONE TAB ONLY
# ======================================================

import webbrowser
import threading

_browser_launched = False

def launch_browser_once():
    global _browser_launched
    if _browser_launched:
        return
    _browser_launched = True

    def open_browser():
        webbrowser.open("http://127.0.0.1:5050/search-hub")

    threading.Timer(1.0, open_browser).start()

# ======================================================
# ROUTES
# ======================================================

@app.route("/")
def index():
    recent = Recent.query.order_by(Recent.opened_at.desc()).limit(10).all()
    paths = [r.file_path for r in recent]
    pinned_paths = set()
    if paths:
        pinned_paths = {
            ref.file_path
            for ref in Reference.query.filter(
                Reference.file_path.in_(paths),
                Reference.pinned == True
            ).all()
        }
    return render_template("index.html", recent=recent, pinned_paths=pinned_paths)
@app.route("/ping")
def ping():
    print("🔥 PING HIT")
    return "pong"
@app.route("/search", methods=["POST"])
def search():
    data = request.json or {}
    query = data.get("query", "").strip()

    if not query:
        return jsonify({
            "results": [],
            "examples": [],
            "fallback_message": None
        })

    mode = "filename"

    try:
        results = run_search(query, mode)
        examples = load_excel_examples(query)

        print("🔥 RESULTS:", len(results))
        print("🔥 EXAMPLES:", len(examples))

        return jsonify({
            "results": results,
            "examples": examples,
            "fallback_message": None if results else "No results found in Q:\\"
        })

    except Exception as e:
        print("🔥 SEARCH CRASH:", repr(e))
        return jsonify({
            "results": [],
            "examples": [],
            "fallback_message": "Search failed."
        }), 500
def is_index_empty():
    try:
        return FileIndex.query.first() is None
    except Exception as exc:
        print(f"⚠️ Could not check file index state: {exc}")
        return True
@app.route("/search-hub")
def search_hub():
    first_time = is_index_empty()
    return render_template("search.html", first_time=first_time)


@app.route("/reindex", methods=["POST"])
def reindex():
    try:
        # 1. Clear existing index
        db.session.execute(text("DELETE FROM file_index"))
        db.session.commit()

        # 2. Rebuild the index using your existing index function
        index_qdrive_files()

        return jsonify({"success": True})
    except Exception as e:
        print("❌ Reindex error:", e)
        return jsonify({"success": False, "error": str(e)})



@app.route("/open-recent/<int:recent_id>")
def open_recent_file(recent_id):
    item = Recent.query.get_or_404(recent_id)
    normalized = os.path.normpath(item.file_path)
    try:
        os.startfile(normalized)
    except Exception as e:
        return f"Error opening file: {e}", 500
    return redirect(url_for("index"))



# ======================================================
# REFERENCE LIBRARY (WITH BUILT-IN GROUPS)
# ======================================================

BUILT_IN_GROUPS = [
    "Ungrouped",
    "Procedures",
    "Cut Sheets",
    "Tracts",
    "Parcel Maps",
    "General"
]
@app.route("/reference/update/<int:ref_id>", methods=["POST"])
def update_reference_field(ref_id):
    data = request.get_json() or {}

    ref = Reference.query.get(ref_id)

    if not ref:
        return jsonify({
            "success": False,
            "message": "Reference not found."
        }), 404

    field = data.get("field")
    value = data.get("value", "")

    allowed_fields = {
        "display_name",
        "group",
        "notes",
    }

    if field not in allowed_fields:
        return jsonify({
            "success": False,
            "message": "Invalid field."
        }), 400

    try:
        setattr(ref, field, value)
        db.session.commit()

        return jsonify({
            "success": True,
            "ref_id": ref.id,
            "field": field,
            "value": value
        }), 200

    except SQLAlchemyError as e:
        db.session.rollback()
        print("💥 update_reference_field ERROR:", e)

        return jsonify({
            "success": False,
            "message": str(e)
        }), 500
@app.route("/copy-to-workspace", methods=["POST"])
def copy_to_workspace():
    data = request.get_json() or {}

    file_path = data.get("file_path", "").strip()
    group_name = data.get("group", "Ungrouped").strip()

    result = copy_reference_to_workspace(
        source_path=file_path,
        group_name=group_name
    )

    print("📥 COPY TO WORKSPACE:", result)

    status_code = 200 if result.get("success") else 400

    return jsonify(result), status_code

@app.route("/reference")
def reference_library():
    # 1. Load pinned refs
    refs = (
        Reference.query.filter_by(pinned=True)
        .order_by(Reference.added_at.desc())
        .all()
    )

    # 2. Group items
    grouped = defaultdict(list)
    for r in refs:
        group_name = r.group if r.group else "Ungrouped"
        grouped[group_name].append(r)

    # 3. Load custom groups from DB
    groups = CustomGroup.query.order_by(CustomGroup.name.asc()).all()

    # 4. Build map: group name → ID
    group_ids = {g.name: g.id for g in groups}

    # 5. Built-in group names
    built_in_groups = [
        "Ungrouped",
        "Procedures",
        "Parcel Maps",
        "Cut Sheets",
        "Tracts",
        "General",
    ]

    # 6. Render template
    return render_template(
        "reference.html",
        grouped=grouped,
        custom_groups=[g.name for g in groups],
        built_in_groups=built_in_groups,
        group_ids=group_ids,
    )


# ======================================================
# CREATE CUSTOM GROUP
# ======================================================
@app.route("/create-group", methods=["POST"])
def create_group():
    data = request.get_json() or {}
    name = data.get("name", "").strip()

    if not name:
        return jsonify({"success": False, "error": "Empty group name"}), 400

    # Block duplicates including trying to recreate built-ins
    if name in BUILT_IN_GROUPS:
        return jsonify({"success": False, "error": "Group already exists"}), 400

    existing = CustomGroup.query.filter_by(name=name).first()
    if existing:
        return jsonify({"success": False, "error": "Group already exists"}), 400

    g = CustomGroup(name=name)
    db.session.add(g)
    db.session.commit()

    return jsonify({"success": True, "id": g.id, "name": g.name})

from urllib.parse import unquote

from urllib.parse import unquote

@app.route("/delete-group/<path:group_name>", methods=["POST"])
def delete_group(group_name):
    name = unquote(group_name)

    built_in = ["Ungrouped", "Procedures", "Parcel Maps", "Cut Sheets", "Tracts", "General"]
    if name in built_in:
        return jsonify({"error": "cannot delete built-in group"}), 400

    try:
        # Move any references in this group back to Ungrouped
        refs = Reference.query.filter_by(group=name).all()
        for r in refs:
            r.group = "Ungrouped"

        # Delete the CustomGroup row, if it exists
        cg = CustomGroup.query.filter_by(name=name).first()
        if cg:
            db.session.delete(cg)

        db.session.commit()
        return jsonify({"success": True})

    except Exception as e:
        db.session.rollback()
        print("❌ Error deleting group:", e)
        return jsonify({"error": "server error"}), 500


@app.route("/open-folder")
def open_folder():
    path = request.args.get("path")
    if not path:
        return jsonify({"status": "error"}), 400

    try:
        os.system(f'explorer.exe /select,"{path}"')
        return jsonify({"status": "ok"})
    except Exception as e:
        return jsonify({"status": "error", "message": str(e)}), 500


@app.route("/save-display-name/<int:ref_id>", methods=["POST"])
def save_display_name(ref_id):
    data = request.get_json()
    new_name = data.get("display_name", "").strip()

    ref = Reference.query.get_or_404(ref_id)
    ref.display_name = new_name
    db.session.commit()

    return jsonify({"success": True})


@app.route("/save-group/<int:ref_id>", methods=["POST"])
def save_group(ref_id):
    data = request.get_json()
    new_group = data.get("group", "").strip()

    ref = Reference.query.get_or_404(ref_id)
    ref.group = new_group
    db.session.commit()

    return jsonify({"success": True})


@app.route("/update-reference/<int:ref_id>", methods=["POST"])
def update_reference(ref_id):
    ref = Reference.query.get(ref_id)
    if not ref:
        return jsonify({"success": False, "error": "Reference not found"}), 404

    data = request.get_json() or {}

    if "display_name" in data:
        ref.display_name = data["display_name"]
    if "group" in data:
        ref.group = data["group"]
    if "notes" in data:
        ref.notes = data["notes"]

    db.session.commit()
    return jsonify({"success": True})


@app.route("/open-file")
def open_file():
    path = request.args.get("path")
    if not path:
        return "", 400

    normalized = os.path.normpath(path)

    entry = Recent(
        file_name=os.path.basename(normalized),
        file_path=normalized,
        folder=os.path.dirname(normalized),
        file_type=os.path.splitext(normalized)[1],
    )
    db.session.add(entry)
    db.session.commit()

    try:
        os.startfile(normalized)
    except Exception as e:
        print("Error opening file:", e)
        return f"Error opening file: {e}", 500

    return "", 204


# ======================================================
# EXPORT SETTINGS (FINAL)
# ======================================================
from sqlalchemy import text

@app.route("/export-settings", methods=["GET"])
def export_settings():
    rows = Reference.query.filter_by(pinned=True).all()

    items = []
    for r in rows:
        items.append({
            "id": r.id,
            "file_name": r.file_name,
            "file_path": r.file_path,
            "folder": r.folder,
            "display_name": r.display_name,
            "notes": r.notes,
            "group": r.group,
            "pinned": r.pinned
        })

    groups = [g.name for g in CustomGroup.query.all()]

    return jsonify({
        "items": items,
        "custom_groups": groups
    })


@app.route("/import-settings", methods=["POST"])
def import_settings():
    data = request.json
    groups = data.get("custom_groups", [])
    items = data.get("items", [])

    # Clear old data
    db.session.execute(text("DELETE FROM custom_group"))
    db.session.execute(text("DELETE FROM reference"))

    # Insert groups
    for g in groups:
        db.session.execute(
            text("INSERT INTO custom_group (name) VALUES (:name)"),
            {"name": g}
        )

    # Insert items
    for item in items:
        db.session.execute(text("""
            INSERT INTO reference
            (file_name, file_path, folder, display_name, notes, "group", pinned)
            VALUES (:file_name, :file_path, :folder, :display_name, :notes, :group, :pinned)
        """), {
            "file_name": item["file_name"],
            "file_path": item["file_path"],
            "folder": item.get("folder"),
            "display_name": item.get("display_name"),
            "notes": item.get("notes"),
            "group": item.get("group"),
            "pinned": int(item.get("pinned", 0))
        })

    db.session.commit()
    return jsonify({"success": True, "message": "Settings imported successfully!"})


if __name__ == "__main__":
    app.run(
        host="127.0.0.1",
        port=5050,
        debug=True,
        use_reloader=False
    )