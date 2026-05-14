import os
import re

# ======================================================
# PATH GUARD — matches existing service pattern verbatim
# ======================================================

ALLOWED_SOURCE_ROOTS = [
    r"Q:\WORD",
    r"Q:\Excel Files",
    r"Q:\Mapping Reference Workspace",
]


def _normalize_path(path):
    return os.path.normcase(os.path.abspath(os.path.normpath(path)))


def _is_allowed_source_path(source_path):
    normalized_source = _normalize_path(source_path)

    for root in ALLOWED_SOURCE_ROOTS:
        normalized_root = _normalize_path(root)

        try:
            if os.path.commonpath([normalized_source, normalized_root]) == normalized_root:
                return True
        except ValueError:
            continue

    return False


# ======================================================
# EXTRACTION CAPS
# ======================================================

MAX_EXCEL_ROWS = 10
MAX_EXCEL_SHEETS = 3
MAX_DOCX_PARAGRAPHS = 30
MAX_PDF_PAGES = 5
MAX_APNS = 20
MAX_DRNS = 20
MAX_DATES = 15


# ======================================================
# REGEX — APNS
# ======================================================
# OC Assessor formats: 123-456-78  and  123-45-678

APN_PATTERN = re.compile(
    r"\b(\d{3}-\d{3}-\d{2}|\d{3}-\d{2}-\d{3})\b"
)


# ======================================================
# REGEX — DRNS (context-anchored only)
# ======================================================
# Only extract numbers that follow a recognized label.
# Never treat every 7-9 digit number as a DRN.

DRN_LABEL_PATTERN = re.compile(
    r"(?:DRN|Doc(?:ument)?\.?\s*No\.?|Document\s*#|Recording\s*No\.?"
    r"|Instrument\s*No\.?|Recorded\s+as\s+Doc(?:ument)?\.?\s*No\.?)"
    r"[\s:#\-]*"
    r"(\d{4}[-\s]?\d{4,7}|\d{6,9})",
    re.IGNORECASE,
)


# ======================================================
# REGEX — DATES
# ======================================================

_DATE_PATTERNS = [
    # MM/DD/YYYY or MM-DD-YYYY
    re.compile(
        r"\b(0?[1-9]|1[0-2])[/\-](0?[1-9]|[12]\d|3[01])[/\-](19|20)\d{2}\b"
    ),
    # Month DD, YYYY (full name)
    re.compile(
        r"\b(?:January|February|March|April|May|June|July|August"
        r"|September|October|November|December)\s+\d{1,2},?\s+(?:19|20)\d{2}\b",
        re.IGNORECASE,
    ),
    # Mon. DD, YYYY (abbreviated)
    re.compile(
        r"\b(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)\.?\s+"
        r"\d{1,2},?\s+(?:19|20)\d{2}\b",
        re.IGNORECASE,
    ),
]


# ======================================================
# MAPPING KEYWORDS
# ======================================================

MAPPING_KEYWORDS = [
    # Map / instrument types
    "Parcel Map", "Parcel Merger", "Parcel Split", "Parcel Consolidation",
    "Tract Map", "Tentative Map", "Final Map", "Subdivision Map",
    "Record of Survey", "Certificate of Compliance",
    "Lot Line Adjustment", "LLA", "Notice of Lot Line Adjustment",
    # Deed / conveyance
    "Grant Deed", "Quitclaim Deed", "Trust Deed", "Deed of Trust",
    "Affidavit", "Declaration", "Certificate", "Resolution",
    "Notice of Completion",
    # Easement / R-O-W
    "Easement", "Right of Way", "Right-of-Way", "R/W",
    "Access Easement", "Drainage Easement", "Utility Easement",
    "Slope Easement", "Dedication", "Vacation",
    # Land restructuring
    "Reversion to Acreage", "Annexation", "Detachment",
    "Boundary Adjustment",
    # Condo / planned development
    "Condominium Plan", "Condo Plan", "CC&R", "Planned Unit Development",
    # Legal description / survey terms
    "Legal Description", "Metes and Bounds",
    "Point of Beginning", "Point of Commencement", "T.P.O.B.",
    "True Point of Beginning", "Thence", "Bearings",
    "Northerly", "Southerly", "Easterly", "Westerly",
    "Monument", "Benchmark", "Control Point",
    "Section", "Township", "Range", "Meridian",
    "Lot", "Block", "Tract",
    # R/T codes
    "Revenue and Taxation Code", "R&T", "R/T Code",
    # OC Assessor / ATS internal
    "ATS", "ArcGIS", "DOC Screening", "Parceling Log",
    "Type A", "Type B", "Type C", "Type D",
    "Type E", "Type F", "Type G",
    # Parcel rules / notes
    "Parent Parcel", "Child Parcel", "Resulting Parcel",
    "Merger", "Split", "Remnant",
    "Tax Default", "Redemption",
    # Reference indicators
    "Assessor's Parcel", "APN",
    "Address Assignment", "Zone Change",
    "General Plan", "Setback", "Building Line",
]


# ======================================================
# CLASSIFICATION SIGNALS
# ======================================================

_CLASSIFICATION_SIGNALS = [
    ("Parcel Map",                  ["parcel map"]),
    ("Lot Line Adjustment",         ["lot line adjustment", " lla ", "notice of lot line"]),
    ("Tract Map / Subdivision",     ["tract map", "final map", "tentative map", "subdivision"]),
    ("Easement / Right-of-Way",     ["easement", "right of way", "right-of-way", " r/w ", "dedication", "vacation"]),
    ("Deed",                        ["grant deed", "quitclaim deed", "trust deed", "deed of trust"]),
    ("Condominium Plan",            ["condominium plan", "condo plan", "cc&r"]),
    ("Notice of Completion",        ["notice of completion"]),
    ("Certificate of Compliance",   ["certificate of compliance"]),
    ("Legal Description / Survey",  ["metes and bounds", "legal description", "record of survey",
                                     "point of beginning", "thence"]),
    ("Parcel Merger",               ["parcel merger", "merger of parcels", "parcel consolidation"]),
    ("SOP / Procedure Guide",       ["standard operating", " sop ", "procedure steps", "step-by-step"]),
    ("Reversion to Acreage",        ["reversion to acreage"]),
    ("Annexation",                  ["annexation", "detachment"]),
]


# ======================================================
# EXTRACTION — EXCEL
# ======================================================

def _extract_excel(file_path):
    import pandas as pd

    sections = []
    preview_lines = []
    text_parts = []

    xl = pd.ExcelFile(file_path)
    sections = list(xl.sheet_names)

    for sheet_name in xl.sheet_names[:MAX_EXCEL_SHEETS]:
        df = xl.parse(sheet_name, nrows=MAX_EXCEL_ROWS, dtype=str)
        df = df.fillna("")

        preview_lines.append(f"[Sheet: {sheet_name}]")

        col_names = [
            str(c) for c in df.columns
            if str(c).strip() and not str(c).startswith("Unnamed")
        ]
        if col_names:
            header_line = "Columns: " + " | ".join(col_names[:10])
            preview_lines.append(header_line)
            text_parts.append(" ".join(col_names))

        for _, row in df.head(MAX_EXCEL_ROWS).iterrows():
            cells = [str(v).strip() for v in row.values if str(v).strip()]
            if cells:
                row_line = " | ".join(cells)
                preview_lines.append(row_line)
                text_parts.append(row_line)

    return "\n".join(text_parts), sections, preview_lines[:30]


# ======================================================
# EXTRACTION — DOCX
# ======================================================

def _extract_docx(file_path):
    from docx import Document

    sections = []
    preview_lines = []
    text_parts = []
    para_count = 0

    doc = Document(file_path)

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        text_parts.append(text)

        if para_count < MAX_DOCX_PARAGRAPHS:
            preview_lines.append(text)
            para_count += 1

        if para.style and "Heading" in para.style.name:
            sections.append(text)

    for table in doc.tables[:2]:
        for row in table.rows[:8]:
            cells = [c.text.strip() for c in row.cells[:6] if c.text.strip()]
            if cells:
                row_line = " | ".join(cells)
                text_parts.append(row_line)
                if len(preview_lines) < MAX_DOCX_PARAGRAPHS + 10:
                    preview_lines.append(row_line)

    return "\n".join(text_parts), sections, preview_lines


# ======================================================
# EXTRACTION — PDF
# ======================================================

def _extract_pdf(file_path):
    from pypdf import PdfReader  # type: ignore

    sections = []
    preview_lines = []
    text_parts = []

    reader = PdfReader(file_path)
    total = len(reader.pages)
    sections = [f"Page {i + 1}" for i in range(min(total, MAX_PDF_PAGES))]

    for page in reader.pages[:MAX_PDF_PAGES]:
        page_text = page.extract_text() or ""
        if not page_text.strip():
            continue

        text_parts.append(page_text)

        for line in page_text.splitlines():
            line = line.strip()
            if line and len(preview_lines) < 25:
                preview_lines.append(line)

    return "\n".join(text_parts), sections, preview_lines


# ======================================================
# ENTITY DETECTION
# ======================================================

def _detect_apns(text):
    found = []
    seen = set()
    for match in APN_PATTERN.finditer(text):
        apn = match.group(0)
        if apn not in seen:
            seen.add(apn)
            found.append(apn)
        if len(found) >= MAX_APNS:
            break
    return found


def _detect_drns(text):
    found = []
    seen = set()
    for match in DRN_LABEL_PATTERN.finditer(text):
        raw = match.group(1)
        digits = re.sub(r"[\s\-]", "", raw)
        if len(digits) >= 6 and digits not in seen:
            seen.add(digits)
            found.append(digits)
        if len(found) >= MAX_DRNS:
            break
    return found


def _detect_dates(text):
    found = []
    seen = set()
    for pattern in _DATE_PATTERNS:
        for match in pattern.finditer(text):
            date_str = match.group(0).strip()
            if date_str not in seen:
                seen.add(date_str)
                found.append(date_str)
            if len(found) >= MAX_DATES:
                return found
    return found


def _detect_mapping_keywords(text):
    text_lower = text.lower()
    found = []
    for kw in MAPPING_KEYWORDS:
        if kw.lower() in text_lower:
            found.append(kw)
    return found


# ======================================================
# CLASSIFICATION
# ======================================================

def _classify_document(text):
    text_lower = text.lower()
    matched = []

    for label, signals in _CLASSIFICATION_SIGNALS:
        for signal in signals:
            if signal.lower() in text_lower:
                matched.append(label)
                break

    if not matched:
        return "General Mapping Document"

    return " / ".join(matched[:3])


# ======================================================
# SUMMARY BUILDER
# ======================================================

def _build_summary(file_type_label, what_it_is, apns, drns, dates):
    parts = []

    if what_it_is and what_it_is != "General Mapping Document":
        parts.append(f"This appears to be a {what_it_is}.")
    else:
        parts.append(f"This is a {file_type_label} from the Mapping reference folders.")

    details = []

    if apns:
        sample = apns[0] if len(apns) == 1 else f"{apns[0]} and {len(apns) - 1} more"
        details.append(f"{len(apns)} APN{'s' if len(apns) > 1 else ''} detected ({sample})")

    if drns:
        details.append(f"{len(drns)} DRN{'s' if len(drns) > 1 else ''} detected")

    if dates:
        details.append(f"{len(dates)} date{'s' if len(dates) > 1 else ''} found")

    if details:
        parts.append(", ".join(details) + ".")

    return " ".join(parts)


# ======================================================
# PUBLIC ENTRY POINT
# ======================================================

def summarize_file(file_path):
    """
    Locally extract, classify, and detect entities from a Q:\\ file.
    No cloud calls, no AI, no OCR, no file copies, no source modification.
    """
    if not file_path:
        return {"ok": False, "error": "Missing file path."}

    file_path = os.path.normpath(file_path)

    if not os.path.exists(file_path):
        return {"ok": False, "error": "File not found."}

    if not os.path.isfile(file_path):
        return {"ok": False, "error": "Path is not a file."}

    if not _is_allowed_source_path(file_path):
        return {
            "ok": False,
            "error": "File is outside approved Mapping reference folders.",
        }

    file_name = os.path.basename(file_path)
    ext = os.path.splitext(file_name)[1].lower()

    # --------------------------------------------------
    # Unsupported: images
    # --------------------------------------------------
    if ext in {".png", ".jpg", ".jpeg", ".gif", ".bmp", ".tiff", ".tif"}:
        return {
            "ok": True,
            "supported": False,
            "file_name": file_name,
            "file_type": ext.lstrip(".").upper(),
            "source_path": file_path,
            "error": "Image files cannot be summarized locally. No OCR available.",
        }

    # --------------------------------------------------
    # Unsupported: legacy .doc
    # --------------------------------------------------
    if ext == ".doc":
        return {
            "ok": True,
            "supported": False,
            "file_name": file_name,
            "file_type": "DOC",
            "source_path": file_path,
            "error": "Legacy .doc files are not supported. Open the original file directly.",
        }

    # --------------------------------------------------
    # Unsupported: everything else
    # --------------------------------------------------
    if ext not in {".pdf", ".docx", ".docm", ".xlsx", ".xls"}:
        return {
            "ok": True,
            "supported": False,
            "file_name": file_name,
            "file_type": ext.lstrip(".").upper() if ext else "Unknown",
            "source_path": file_path,
            "error": f"File type '{ext}' is not supported for local summary.",
        }

    # --------------------------------------------------
    # Extract
    # --------------------------------------------------
    try:
        if ext in {".xlsx", ".xls"}:
            raw_text, sections, preview = _extract_excel(file_path)
            file_type_label = "Excel Workbook"

        elif ext in {".docx", ".docm"}:
            raw_text, sections, preview = _extract_docx(file_path)
            file_type_label = "Word Macro-Enabled Document" if ext == ".docm" else "Word Document"

        else:  # .pdf
            raw_text, sections, preview = _extract_pdf(file_path)
            file_type_label = "PDF"

    except PermissionError as exc:
        return {"ok": False, "error": str(exc)}

    except Exception as exc:
        return {"ok": False, "error": f"Could not read file: {exc}"}

    # --------------------------------------------------
    # Scanned / image-only PDF guard
    # --------------------------------------------------
    if ext == ".pdf" and not raw_text.strip():
        return {
            "ok": True,
            "supported": False,
            "file_name": file_name,
            "file_type": "PDF",
            "source_path": file_path,
            "error": (
                "PDF contains no readable text. "
                "It may be a scanned or image-based document."
            ),
        }

    # --------------------------------------------------
    # Entity detection
    # --------------------------------------------------
    apns = _detect_apns(raw_text)
    drns = _detect_drns(raw_text)
    dates = _detect_dates(raw_text)
    keywords = _detect_mapping_keywords(raw_text)

    # --------------------------------------------------
    # Classification + summary
    # --------------------------------------------------
    what_it_is = _classify_document(raw_text)
    summary = _build_summary(file_type_label, what_it_is, apns, drns, dates)

    return {
        "ok": True,
        "supported": True,
        "file_name": file_name,
        "file_type": file_type_label,
        "source_path": file_path,
        "detected_sections": sections,
        "summary": summary,
        "preview_lines": preview,
        "possible_apns": apns,
        "possible_drns": drns,
        "possible_dates": dates,
        "possible_keywords": keywords,
        "what_this_appears_to_be": what_it_is,
    }