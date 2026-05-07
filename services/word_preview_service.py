import os
from docx import Document


MAX_PARAGRAPHS = 12
MAX_TABLE_ROWS = 8
MAX_TABLE_COLUMNS = 6

ALLOWED_SOURCE_ROOTS = [
    r"Q:\WORD",
    r"Q:\Excel Files",
    r"Q:\Mapping Reference Workspace",
]


def _normalize_path(path):
    return os.path.normcase(os.path.abspath(os.path.normpath(path)))


def _is_allowed_source_path(source_path):
    normalized_source = _normalize_path(source_path)

    for root in ALLOWED_SOURCE_ROOTS:
        normalized_root = _normalize_path(root)

        try:
            if os.path.commonpath([normalized_source, normalized_root]) == normalized_root:
                return True
        except ValueError:
            continue

    return False


def build_word_preview(file_path):
    if not file_path:
        return {
            "success": False,
            "message": "Missing file path.",
        }

    file_path = os.path.normpath(file_path)

    if not os.path.exists(file_path):
        return {
            "success": False,
            "message": "Word file not found.",
            "file_path": file_path,
        }

    if not os.path.isfile(file_path):
        return {
            "success": False,
            "message": "Path is not a file.",
            "file_path": file_path,
        }

    if not _is_allowed_source_path(file_path):
        return {
            "success": False,
            "message": "File is outside approved Mapping reference folders.",
            "file_path": file_path,
        }

    ext = os.path.splitext(file_path)[1].lower()

    if ext == ".doc":
        return {
            "success": False,
            "message": "Legacy .doc preview is not supported. Open the original file.",
            "file_path": file_path,
        }

    if ext != ".docx":
        return {
            "success": False,
            "message": "File is not a supported Word document.",
            "file_path": file_path,
        }

    try:
        document = Document(file_path)

        paragraphs = []
        for paragraph in document.paragraphs:
            text = paragraph.text.strip()

            if text:
                paragraphs.append(text)

            if len(paragraphs) >= MAX_PARAGRAPHS:
                break

        tables = []
        for table in document.tables[:1]:
            table_rows = []

            for row in table.rows[:MAX_TABLE_ROWS]:
                cells = []

                for cell in row.cells[:MAX_TABLE_COLUMNS]:
                    cells.append(cell.text.strip())

                table_rows.append(cells)

            tables.append(table_rows)

        return {
            "success": True,
            "file_path": file_path,
            "file_name": os.path.basename(file_path),
            "paragraphs": paragraphs,
            "tables": tables,
        }

    except PermissionError:
        return {
            "success": False,
            "message": "Word file is currently open or locked. Close the file and try again.",
            "file_path": file_path,
        }

    except Exception as e:
        return {
            "success": False,
            "message": f"Could not preview Word file: {e}",
            "file_path": file_path,
        }